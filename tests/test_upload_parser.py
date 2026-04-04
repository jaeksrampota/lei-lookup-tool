"""Tests for the upload parser."""

import io
from pathlib import Path

import pytest

from src.upload_parser import parse_upload, MAX_ENTITIES

FIXTURES_DIR = Path(__file__).parent / "fixtures"


class TestParseXlsx:
    def test_parse_xlsx_standard(self, sample_xlsx):
        content = sample_xlsx.read_bytes()
        entities = parse_upload("test.xlsx", content)
        assert len(entities) == 3
        assert entities[0].name == "CAIAC Fund Management AG"
        assert entities[2].country == "Česká republika"

    def test_parse_xlsx_skips_empty_rows(self):
        """Create an xlsx with an empty name row - should be skipped."""
        import pandas as pd
        buf = io.BytesIO()
        df = pd.DataFrame({
            "Name": ["Entity A", "", "Entity B"],
            "ISIN": ["", "", ""],
            "Street": ["", "", ""],
            "Town": ["", "", ""],
            "Country": ["", "", ""],
            "ZIP": ["", "", ""],
        })
        df.to_excel(buf, index=False, engine="openpyxl")
        entities = parse_upload("test.xlsx", buf.getvalue())
        assert len(entities) == 2


class TestParseCsv:
    def test_parse_csv_comma(self, sample_csv):
        content = sample_csv.read_bytes()
        entities = parse_upload("test.csv", content)
        assert len(entities) == 3
        assert entities[0].name == "CAIAC Fund Management AG"

    def test_parse_csv_semicolon(self, sample_semicolon_csv):
        content = sample_semicolon_csv.read_bytes()
        entities = parse_upload("test.csv", content)
        assert len(entities) == 3

    def test_parse_csv_encoding_utf8_bom(self):
        content = b'\xef\xbb\xbfName,ISIN,Street,Town,Country,ZIP\nTest Corp,,,,CZ,\n'
        entities = parse_upload("test.csv", content)
        assert len(entities) == 1
        assert entities[0].name == "Test Corp"

    def test_parse_csv_encoding_latin1(self):
        text = "Name,ISIN,Street,Town,Country,ZIP\nCaf\xe9 Corp,,,,FR,\n"
        content = text.encode("latin-1")
        entities = parse_upload("test.csv", content)
        assert len(entities) == 1

    def test_parse_csv_encoding_cp1250(self):
        text = "Name,ISIN,Street,Town,Country,ZIP\nPraha a.s.,,Národní,,Česká republika,\n"
        content = text.encode("cp1250")
        entities = parse_upload("test.csv", content)
        assert len(entities) == 1

    def test_parse_csv_quoted_fields(self):
        content = b'Name,ISIN,Street,Town,Country,ZIP\n"Entity, Inc.",,,,US,\n'
        entities = parse_upload("test.csv", content)
        assert len(entities) == 1
        assert entities[0].name == "Entity, Inc."


class TestParseDocx:
    def test_parse_docx_with_table(self, sample_docx):
        content = sample_docx.read_bytes()
        entities = parse_upload("test.docx", content)
        assert len(entities) >= 3
        names = [e.name for e in entities]
        assert "CAIAC Fund Management AG" in names

    def test_parse_docx_no_table(self):
        """A docx with paragraphs but no tables should parse names from paragraphs."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Alpha Corp")
        doc.add_paragraph("Beta LLC")
        doc.add_paragraph("")  # empty paragraph, should be skipped
        buf = io.BytesIO()
        doc.save(buf)
        entities = parse_upload("test.docx", buf.getvalue())
        assert len(entities) == 2
        assert entities[0].name == "Alpha Corp"


class TestValidation:
    def test_empty_file(self):
        with pytest.raises(ValueError, match="empty"):
            parse_upload("test.csv", b"")

    def test_unsupported_extension(self):
        with pytest.raises(ValueError, match="Unsupported"):
            parse_upload("test.txt", b"some content")

    def test_unsupported_pdf(self):
        with pytest.raises(ValueError, match="Unsupported"):
            parse_upload("test.pdf", b"%PDF-1.4")

    def test_max_entities_enforced(self):
        lines = ["Name,ISIN,Street,Town,Country,ZIP"]
        for i in range(MAX_ENTITIES + 10):
            lines.append(f"Entity {i},,,,CZ,")
        content = "\n".join(lines).encode("utf-8")
        with pytest.raises(ValueError, match="Too many"):
            parse_upload("test.csv", content)

    def test_unicode_preserved(self):
        content = "Name,ISIN,Street,Town,Country,ZIP\nŽelezárny Prostějov a.s.,,Národní třída,,Česká republika,\n"
        entities = parse_upload("test.csv", content.encode("utf-8"))
        assert entities[0].name == "Železárny Prostějov a.s."
