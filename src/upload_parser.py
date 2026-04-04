"""Parse uploaded files (CSV, XLSX, DOCX) into InputEntity lists."""

import io
import logging
from pathlib import Path
from typing import Optional

import pandas as pd

from .models import InputEntity

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_ENTITIES = 500


def parse_upload(filename: str, content: bytes) -> list[InputEntity]:
    """Parse an uploaded file into a list of InputEntity objects.

    Supports: .xlsx, .csv, .docx
    Raises ValueError for unsupported formats or invalid data.
    """
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"File too large ({len(content)} bytes). Maximum is {MAX_FILE_SIZE // (1024*1024)} MB.")

    if not content or len(content) == 0:
        raise ValueError("File is empty.")

    ext = Path(filename).suffix.lower()

    if ext == ".xlsx":
        entities = _parse_xlsx(content)
    elif ext == ".csv":
        entities = _parse_csv(content)
    elif ext == ".docx":
        entities = _parse_docx(content)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported formats: .xlsx, .csv, .docx")

    if not entities:
        raise ValueError("No valid entities found in file. Ensure at least one row has a non-empty name.")

    if len(entities) > MAX_ENTITIES:
        raise ValueError(f"Too many entities ({len(entities)}). Maximum is {MAX_ENTITIES}.")

    logger.info("Parsed %d entities from %s", len(entities), filename)
    return entities


def _parse_xlsx(content: bytes) -> list[InputEntity]:
    """Parse Excel file using same logic as batch.py."""
    buf = io.BytesIO(content)
    df = pd.read_excel(buf, engine="openpyxl")
    return _dataframe_to_entities(df)


def _parse_csv(content: bytes) -> list[InputEntity]:
    """Parse CSV with auto-detected delimiter and encoding."""
    text = _decode_content(content)
    delimiter = _detect_delimiter(text)
    buf = io.StringIO(text)
    df = pd.read_csv(buf, sep=delimiter, dtype=str, keep_default_na=False)
    return _dataframe_to_entities(df)


def _parse_docx(content: bytes) -> list[InputEntity]:
    """Parse Word document tables into entities."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is required to parse .docx files. Install it with: pip install python-docx")

    buf = io.BytesIO(content)
    doc = Document(buf)

    entities: list[InputEntity] = []

    # Try tables first
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # Use first row as header hints
        headers = [cell.text.strip().lower() for cell in rows[0].cells]
        col_map = _map_docx_headers(headers)

        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            entity = _row_to_entity_mapped(cells, col_map)
            if entity and entity.name:
                entities.append(entity)

    # Fallback: one entity name per paragraph
    if not entities:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                entities.append(InputEntity(name=text))

    return entities


def _decode_content(content: bytes) -> str:
    """Try multiple encodings to decode CSV content."""
    for encoding in ("utf-8-sig", "utf-8", "cp1250", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return content.decode("latin-1")  # latin-1 never fails


def _detect_delimiter(text: str) -> str:
    """Auto-detect CSV delimiter (comma vs semicolon)."""
    first_lines = text.split("\n", 5)[:5]
    sample = "\n".join(first_lines)
    semicolons = sample.count(";")
    commas = sample.count(",")
    return ";" if semicolons > commas else ","


def _dataframe_to_entities(df: pd.DataFrame) -> list[InputEntity]:
    """Convert a DataFrame (first 6 columns) to InputEntity list."""
    # Use first 6 columns positionally
    ncols = min(len(df.columns), 6)
    col_names = list(df.columns[:ncols])
    df = df[col_names]

    entities = []
    for _, row in df.iterrows():
        def _get(idx: int) -> Optional[str]:
            if idx < ncols:
                val = row.iloc[idx]
                if pd.notna(val):
                    s = str(val).strip()
                    return s if s else None
            return None

        name = _get(0) or ""
        if not name:
            continue

        entities.append(InputEntity(
            name=name,
            isin=_get(1),
            street=_get(2),
            town=_get(3),
            country=_get(4),
            zip_code=_get(5),
        ))

    return entities


def _map_docx_headers(headers: list[str]) -> dict[str, int]:
    """Map Word table header names to column indices."""
    mapping: dict[str, int] = {}
    for i, h in enumerate(headers):
        h_lower = h.lower().strip()
        if any(k in h_lower for k in ("name", "název", "entity", "firma", "company")):
            mapping.setdefault("name", i)
        elif any(k in h_lower for k in ("isin",)):
            mapping.setdefault("isin", i)
        elif any(k in h_lower for k in ("street", "ulice", "adresa", "address")):
            mapping.setdefault("street", i)
        elif any(k in h_lower for k in ("town", "city", "město", "obec")):
            mapping.setdefault("town", i)
        elif any(k in h_lower for k in ("country", "země", "stát")):
            mapping.setdefault("country", i)
        elif any(k in h_lower for k in ("zip", "psč", "postal")):
            mapping.setdefault("zip_code", i)

    # Fallback: positional if no headers matched
    if "name" not in mapping and len(headers) >= 1:
        mapping = {"name": 0}
        if len(headers) >= 2:
            mapping["isin"] = 1
        if len(headers) >= 3:
            mapping["street"] = 2
        if len(headers) >= 4:
            mapping["town"] = 3
        if len(headers) >= 5:
            mapping["country"] = 4
        if len(headers) >= 6:
            mapping["zip_code"] = 5

    return mapping


def _row_to_entity_mapped(cells: list[str], col_map: dict[str, int]) -> Optional[InputEntity]:
    """Create InputEntity from a row using a column mapping."""
    def _get(field: str) -> Optional[str]:
        idx = col_map.get(field)
        if idx is not None and idx < len(cells):
            val = cells[idx].strip()
            return val if val else None
        return None

    name = _get("name")
    if not name:
        return None

    return InputEntity(
        name=name,
        isin=_get("isin"),
        street=_get("street"),
        town=_get("town"),
        country=_get("country"),
        zip_code=_get("zip_code"),
    )
